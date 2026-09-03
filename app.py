from flask import Flask, render_template, request, send_file
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.colors import HexColor

import csv
import os
import secrets
import tempfile
import time
import zipfile
from datetime import datetime

import pdfplumber
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from tools.lunch_review import extract_employee_page, hours_between
from tools.attendance_review import build_daily_review
from tools.attendance_pdf import create_attendance_pdf
from tools.ad347_data import build_ad347_record
from tools.pdf_generator import create_ad347_pdf
from tools.talent_lms import (
    course_month_end,
    filter_eligible,
    is_completed,
    match_employees,
    read_training_matrix,
    read_ukg_employees,
)
from tools.evaluation_finder import (
    create_evaluation_pdf,
    evaluations_for_month,
    read_evaluation_employees,
)
from tools.birthday_anniversary import (
    create_birthday_anniversary_pdf,
    read_birthdays_and_anniversaries,
)


app = Flask(__name__)

TEMPLATE_PATH = "BLANK_AD347.pdf"
TALENT_UPLOADS = {}
TALENT_UPLOAD_TTL_SECONDS = 60 * 60


@app.route("/attendance-review", methods=["GET", "POST"])
def attendance_review():
    if request.method == "GET":
        return render_template("attendance_review.html")

    exception_file = request.files.get("exception_file")
    lunch_file = request.files.get("lunch_file")

    if not exception_file or not lunch_file:
        return "<h2>Error</h2><p>Both Excel reports are required.</p>", 400

    temp_folder = tempfile.mkdtemp()
    exception_path = os.path.join(temp_folder, "exception_report.xlsx")
    lunch_path = os.path.join(temp_folder, "lunch_report.xlsx")
    exception_file.save(exception_path)
    lunch_file.save(lunch_path)

    try:
        report_date, rows, missing_punches = build_daily_review(
            exception_path,
            lunch_path,
        )

        if missing_punches:
            names = "".join(f"<li>{name}</li>" for name in missing_punches)
            return f"""
                <h2>Fix missing punches first</h2>
                <p>The report was stopped because these employees still have a missing punch:</p>
                <ul>{names}</ul>
                <p>Correct them in UKG, rerun both reports, and upload them again.</p>
                <p><a href='/attendance-review'>Try again</a></p>
            """, 400

        output_path = os.path.join(
            temp_folder,
            "Daily_Attendance_Exception_Report.pdf",
        )
        create_attendance_pdf(report_date, rows, output_path)

        return send_file(
            output_path,
            as_attachment=True,
            download_name=(
                f"Daily_Attendance_Exceptions_{report_date.isoformat()}.pdf"
            ),
            mimetype="application/pdf",
        )
    except Exception as error:
        return f"""
            <h2>Error</h2>
            <p>{error}</p>
            <p><a href='/attendance-review'>Try again</a></p>
        """, 400


@app.route("/")
def home():
    return render_template("index.html")


@app.route("/applicant-tracker")
def applicant_tracker():
    return render_template("applicant_tracker.html")


@app.route("/sunshine-aide-tracker")
def sunshine_aide_tracker():
    return render_template("sunshine_aide_tracker.html")


def review_timesheets(pdf_path):
    grouped_records = {}
    preview_rows = []

    with pdfplumber.open(pdf_path) as document:
        for page in document.pages:
            text = page.extract_text() or ""

            name, department, shifts = extract_employee_page(text)

            missed_dates = []

            for work_date, segments in shifts.items():
                total_hours = sum(
                    hours_between(start, end)
                    for start, end in segments
                )

                if total_hours < 7:
                    continue

                lunch_taken = len(segments) >= 2

                preview_rows.append(
                    {
                        "name": name,
                        "department": department,
                        "date": work_date,
                        "hours": total_hours,
                        "lunch_taken": lunch_taken,
                    }
                )

                if not lunch_taken:
                    missed_dates.append(work_date)

            if missed_dates:
                key = (name, department)
                grouped_records.setdefault(key, [])
                grouped_records[key].extend(missed_dates)

    return grouped_records, preview_rows


def build_preview_page(preview_rows):
    html = """
    <h2>Lunch Review Preview</h2>
    <p><b>No forms were created.</b></p>
    """

    if not preview_rows:
        html += "<p>No qualifying shifts were found.</p>"
    else:
        current_employee = None

        for row in preview_rows:
            employee_key = (row["name"], row["department"])

            if employee_key != current_employee:
                html += f"""
                <hr>
                <h3>{row["name"]}</h3>
                <p><b>Department:</b> {row["department"]}</p>
                """
                current_employee = employee_key

            status = (
                "✅ Lunch Taken"
                if row["lunch_taken"]
                else "❌ No Recorded Lunch"
            )

            html += (
                f'{row["date"]} — '
                f'{row["hours"]:.2f} hours — '
                f'{status}<br>'
            )

    html += """
    <br>
    <p><a href="/lunch-review">Run another review</a></p>
    <p><a href="/">Home</a></p>
    """

    return html


def create_supervisor_pdf(grouped_records, output_path):
    pdf = canvas.Canvas(output_path, pagesize=letter)

    page_width, page_height = letter
    y = page_height - 55

    pdf.setTitle("Supervisor Missed Lunch List")
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(50, y, "Supervisor Missed Lunch List")

    y -= 22
    pdf.setFont("Helvetica", 10)
    pdf.drawString(
        50,
        y,
        "Employees with no recorded meal break on qualifying shifts"
    )

    y -= 30

    sorted_records = sorted(
        grouped_records.items(),
        key=lambda item: (item[0][1], item[0][0]),
    )

    current_department = None

    for (name, department), dates in sorted_records:
        unique_dates = sorted(set(dates))

        if department != current_department:
            if y < 120:
                pdf.showPage()
                y = page_height - 55

            pdf.setFont("Helvetica-Bold", 12)
            pdf.drawString(50, y, department)
            y -= 20
            current_department = department

        if y < 100:
            pdf.showPage()
            y = page_height - 55

            pdf.setFont("Helvetica-Bold", 12)
            pdf.drawString(50, y, department)
            y -= 20

        pdf.setFont("Helvetica-Bold", 10)
        pdf.drawString(65, y, name)

        y -= 15
        pdf.setFont("Helvetica", 10)
        pdf.drawString(
            80,
            y,
            "Missed lunch date(s): " + ", ".join(unique_dates)
        )

        y -= 15
        pdf.drawString(
            80,
            y,
            f"Total missed lunches: {len(unique_dates)}"
        )

        y -= 22

    pdf.setFont("Helvetica-Oblique", 8)
    pdf.drawString(
        50,
        35,
        "This list is for supervisor follow-up during the pay period."
    )

    pdf.save()


def create_final_zip(grouped_records, temp_folder):
    pdf_files = []
    summary_rows = []

    sorted_records = sorted(
        grouped_records.items(),
        key=lambda item: item[0][0],
    )

    for index, ((name, department), dates) in enumerate(
        sorted_records,
        start=1,
    ):
        unique_dates = sorted(set(dates))

        record = build_ad347_record(
            name=name,
            department=department,
            dates=unique_dates,
        )

        safe_name = (
            name.replace(" ", "_")
            .replace("/", "-")
            .replace("\\", "-")
        )

        output_path = os.path.join(
            temp_folder,
            f"{index}_{safe_name}_AD347.pdf",
        )

        create_ad347_pdf(
            record=record,
            template_path=TEMPLATE_PATH,
            output_path=output_path,
        )

        pdf_files.append(output_path)

        summary_rows.append(
            {
                "Employee Name": name,
                "Department": department,
                "Missed Lunch Dates": ", ".join(unique_dates),
                "Total Missed Lunches": len(unique_dates),
            }
        )

    summary_path = os.path.join(
        temp_folder,
        "Lunch_Review_Summary.csv",
    )

    with open(
        summary_path,
        "w",
        newline="",
        encoding="utf-8-sig",
    ) as summary_file:
        fieldnames = [
            "Employee Name",
            "Department",
            "Missed Lunch Dates",
            "Total Missed Lunches",
        ]

        writer = csv.DictWriter(
            summary_file,
            fieldnames=fieldnames,
        )

        writer.writeheader()
        writer.writerows(summary_rows)

    zip_path = os.path.join(
        temp_folder,
        "AD347_Forms_and_Summary.zip",
    )

    with zipfile.ZipFile(zip_path, "w") as zip_file:
        zip_file.write(
            summary_path,
            arcname="Lunch_Review_Summary.csv",
        )

        for pdf_file in pdf_files:
            zip_file.write(
                pdf_file,
                arcname=os.path.basename(pdf_file),
            )

    return zip_path


@app.route("/lunch-review", methods=["GET", "POST"])
def lunch_review():
    if request.method == "POST":
        pdf = request.files.get("pdf")
        action = request.form.get("action")

        if not pdf or pdf.filename == "":
            return "No file selected"

        temp_folder = tempfile.mkdtemp()
        temp_pdf = os.path.join(temp_folder, "timesheets.pdf")
        pdf.save(temp_pdf)

        try:
            grouped_records, preview_rows = review_timesheets(temp_pdf)

            if action == "preview":
                return build_preview_page(preview_rows)

            if action == "supervisor":
                if not grouped_records:
                    return """
                    <h2>No missed lunches found.</h2>
                    <p><a href="/lunch-review">Run another review</a></p>
                    """

                supervisor_path = os.path.join(
                    temp_folder,
                    "Supervisor_Missed_Lunch_List.pdf",
                )

                create_supervisor_pdf(
                    grouped_records,
                    supervisor_path,
                )

                return send_file(
                    supervisor_path,
                    as_attachment=True,
                    download_name="Supervisor_Missed_Lunch_List.pdf",
                )

            if action == "final":
                if not grouped_records:
                    return """
                    <h2>No AD-347 forms needed.</h2>
                    <p><a href="/lunch-review">Run another review</a></p>
                    """

                zip_path = create_final_zip(
                    grouped_records,
                    temp_folder,
                )

                return send_file(
                    zip_path,
                    as_attachment=True,
                    download_name="AD347_Forms_and_Summary.zip",
                )

            return "No action selected"

        except Exception as error:
            return f"<h2>Error</h2><pre>{error}</pre>"

    return """
    <h2>Lunch Review</h2>

    <form method="POST" enctype="multipart/form-data">
        <input
            type="file"
            name="pdf"
            accept=".pdf"
            required
        >

        <br><br>

        <button type="submit" name="action" value="preview">
            Preview Lunch Review
        </button>

        <br><br>

        <button type="submit" name="action" value="supervisor">
            Print Supervisor List
        </button>

        <br><br>

        <button type="submit" name="action" value="final">
            Generate Final AD-347 Forms
        </button>
    </form>

    <p>
        Use Preview or Supervisor List during the pay period.
        Use Final only after the pay period is complete.
    </p>

    <p><a href="/">Home</a></p>
    """


def _talent_department(employee):
    department = employee["department"].strip().upper()
    job = employee["job"].strip().upper()
    if department == "NSG" and job == "CNA":
        return "CNAs"
    if department == "NSG":
        return "Nursing"
    return employee["department"] or "No Department"


def _safe_filename(value):
    safe = "".join(character if character.isalnum() else "_" for character in value)
    return safe.strip("_")[:80] or "InService"


def create_talent_lms_docx(follow_up, report_label, department, eligible_count, completed_count, output_path):
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.right_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.1

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run(f"{report_label} - Follow-Up")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(31, 78, 120)

    summary = document.add_paragraph()
    summary.paragraph_format.space_after = Pt(12)
    summary_run = summary.add_run(
        f"Department: {department}    Eligible: {eligible_count}    "
        f"Completed: {completed_count}    Follow-up: {len(follow_up)}"
    )
    summary_run.bold = True

    note = document.add_paragraph("Give this document only to the manager responsible for this group.")
    note.paragraph_format.space_after = Pt(14)

    if not follow_up:
        document.add_paragraph("Everyone in this group has a completed record.")
    else:
        for item in sorted(follow_up, key=lambda row: (row["employee"]["last"], row["employee"]["first"])):
            employee = item["employee"]
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.15)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.add_run(employee["display_name"])
            if item.get("note"):
                note_run = paragraph.add_run(f" — {item['note']}")
                note_run.italic = True
                note_run.font.color.rgb = RGBColor(128, 74, 0)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("For supervisor follow-up. UKG determines active employee eligibility.")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(8)

    document.save(output_path)


def create_talent_name_review_docx(missing, duplicates, report_label, output_path):
    document = Document()
    title = document.add_paragraph()
    run = title.add_run(f"{report_label} - HR Name Review")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 120)
    document.add_paragraph(
        "These employees could not be matched safely. Correct names or enrollment in TalentLMS, "
        "export a new Training Matrix, and rerun the tool."
    )
    if missing:
        heading = document.add_paragraph()
        heading.add_run("Not found in Training Matrix").bold = True
        for employee in sorted(missing, key=lambda row: (row["department"], row["last"], row["first"])):
            document.add_paragraph(
                f"{employee['display_name']} — {_talent_department(employee)} — {employee['job'] or 'No job listed'}",
                style="List Bullet",
            )
    if duplicates:
        heading = document.add_paragraph()
        heading.add_run("Duplicate exact names in Training Matrix").bold = True
        for item in sorted(duplicates, key=lambda row: row["employee"]["display_name"]):
            employee = item["employee"]
            document.add_paragraph(
                f"{employee['display_name']} — {_talent_department(employee)} — "
                f"{len(item['matrix_names'])} TalentLMS rows",
                style="List Bullet",
            )
    if not missing and not duplicates:
        document.add_paragraph("No name-review problems were found.")
    document.save(output_path)


def _talent_upload_directory(token):
    if not token or not token.isalnum() or len(token) != 32:
        raise ValueError("The uploaded-file session is invalid. Upload both reports again.")
    directory = os.path.join(tempfile.gettempdir(), "landmark_talent_uploads", token)
    if not os.path.isdir(directory):
        raise ValueError("The uploaded reports expired. Upload both reports again.")
    return directory


@app.route("/talent-lms", methods=["GET", "POST"])
def talent_lms():
    if request.method == "GET":
        return render_template("talent_lms.html")

    action = request.form.get("action", "upload")
    try:
        if action == "upload":
            uploaded_files = [file for file in request.files.getlist("report_files") if file.filename]
            if len(uploaded_files) != 2 or any(not file.filename.lower().endswith(".xlsx") for file in uploaded_files):
                raise ValueError("Upload exactly two Excel files: the UKG Employee Information report and TalentLMS Training Matrix.")
            token = secrets.token_hex(16)
            upload_directory = os.path.join(tempfile.gettempdir(), "landmark_talent_uploads", token)
            os.makedirs(upload_directory, exist_ok=False)
            paths = []
            for index, uploaded_file in enumerate(uploaded_files):
                path = os.path.join(upload_directory, f"report_{index}.xlsx")
                uploaded_file.save(path)
                paths.append(path)

            ukg_path = matrix_path = None
            employees = courses = matrix_rows = None
            for path in paths:
                try:
                    candidate_courses, candidate_rows = read_training_matrix(path)
                    matrix_path, courses, matrix_rows = path, candidate_courses, candidate_rows
                    continue
                except Exception:
                    pass
                try:
                    candidate_employees = read_ukg_employees(path)
                    ukg_path, employees = path, candidate_employees
                except Exception:
                    pass
            if not ukg_path or not matrix_path:
                raise ValueError("The tool could not identify one UKG Employee Information file and one TalentLMS Training Matrix file.")

            os.replace(ukg_path, os.path.join(upload_directory, "ukg.xlsx"))
            os.replace(matrix_path, os.path.join(upload_directory, "matrix.xlsx"))
            unique, duplicates, missing = match_employees(employees, matrix_rows)
            current_year = datetime.now().year
            visible_courses = []
            for course in courses:
                month_end = course_month_end(course["name"])
                if month_end and month_end.year < current_year:
                    continue
                visible_courses.append(course)
            return render_template(
                "talent_lms_select.html",
                token=token,
                courses=visible_courses,
                departments=sorted({employee["department"] for employee in employees} | {"CNA"}),
                jobs=sorted({employee["job"] for employee in employees if employee["job"]} | {"DTY", "MRD", "MTN", "ACT", "SSV"}),
                employee_count=len(employees),
                matrix_user_count=len(matrix_rows),
                unique_count=len(unique),
                duplicate_count=len(duplicates),
                missing_count=len(missing),
            )

        upload_directory = _talent_upload_directory(request.form.get("token", ""))
        employees = read_ukg_employees(os.path.join(upload_directory, "ukg.xlsx"))
        courses, matrix_rows = read_training_matrix(os.path.join(upload_directory, "matrix.xlsx"))
        course_index = int(request.form.get("course_index", "-1"))
        course = next((item for item in courses if item["index"] == course_index), None)
        if not course:
            raise ValueError("Select an in-service course.")
        scope = request.form.get("scope", "all")
        if scope not in {"all", "departments", "jobs"}:
            raise ValueError("Select who is required to complete this in-service.")
        eligibility_end = course_month_end(course["name"])
        eligible = filter_eligible(
            employees,
            scope,
            request.form.getlist("departments"),
            request.form.getlist("jobs"),
            eligibility_end,
        )
        if not eligible:
            raise ValueError("No UKG employees matched the selected requirement group.")

        unique, duplicates, missing = match_employees(eligible, matrix_rows)
        completed_ids = set()
        follow_up = []
        for employee, matrix_row in unique:
            if is_completed(matrix_row, course_index):
                completed_ids.add(employee["employee_id"])
            else:
                follow_up.append({"employee": employee, "note": ""})
        for employee in missing:
            follow_up.append({"employee": employee, "note": "Not found in TalentLMS — HR review"})
        for item in duplicates:
            follow_up.append({"employee": item["employee"], "note": "Duplicate TalentLMS name — HR review"})

        report_label = request.form.get("report_title", "").strip() or course["name"]
        safe_label = _safe_filename(report_label)
        result_directory = tempfile.mkdtemp()
        zip_path = os.path.join(result_directory, f"TalentLMS_{safe_label}_Department_Lists.zip")
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as archive:
            departments = sorted({_talent_department(employee) for employee in eligible})
            for department in departments:
                department_employees = [employee for employee in eligible if _talent_department(employee) == department]
                employee_ids = {employee["employee_id"] for employee in department_employees}
                department_follow_up = [item for item in follow_up if item["employee"]["employee_id"] in employee_ids]
                docx_name = f"{_safe_filename(department)}_{safe_label}_Follow_Up.docx"
                docx_path = os.path.join(result_directory, docx_name)
                create_talent_lms_docx(
                    department_follow_up,
                    report_label,
                    department,
                    len(department_employees),
                    len(employee_ids.intersection(completed_ids)),
                    docx_path,
                )
                archive.write(docx_path, arcname=docx_name)
            review_path = os.path.join(result_directory, "HR_Name_Review.docx")
            create_talent_name_review_docx(missing, duplicates, report_label, review_path)
            archive.write(review_path, arcname="HR_Name_Review.docx")

        return send_file(zip_path, as_attachment=True, download_name=os.path.basename(zip_path), mimetype="application/zip")
    except Exception as error:
        return f"<h2>Error</h2><p>{error}</p><p><a href='/talent-lms'>Start over</a></p>", 400


@app.route("/ad347")
def ad347():
    return "<h2>Coming Soon</h2>"


@app.route("/evaluations", methods=["GET", "POST"])
def evaluations():
    if request.method == "GET":
        return render_template(
            "evaluations.html",
            default_month=datetime.now().strftime("%Y-%m"),
        )

    ukg_file = request.files.get("ukg_file")
    report_month = request.form.get("report_month", "")
    if not ukg_file or not report_month:
        return "<h2>Error</h2><p>The UKG report and evaluation month are required.</p>", 400

    try:
        selected_month = datetime.strptime(report_month, "%Y-%m")
        temp_folder = tempfile.mkdtemp()
        input_path = os.path.join(temp_folder, "ukg_evaluations.xlsx")
        ukg_file.save(input_path)
        employees = read_evaluation_employees(input_path)
        rows = evaluations_for_month(employees, selected_month.year, selected_month.month)
        if not rows:
            return f"""
                <h2>No evaluations due</h2>
                <p>No 90-day or annual evaluations were found for {selected_month.strftime('%B %Y')}.</p>
                <p><a href='/evaluations'>Choose another month</a></p>
            """

        output_name = f"Department_Evaluations_{report_month}.pdf"
        output_path = os.path.join(temp_folder, output_name)
        create_evaluation_pdf(rows, selected_month.year, selected_month.month, output_path)
        return send_file(
            output_path,
            as_attachment=True,
            download_name=output_name,
            mimetype="application/pdf",
        )
    except Exception as error:
        return f"""
            <h2>Could not create the evaluation packet</h2>
            <p>{error}</p>
            <p><a href='/evaluations'>Try again</a></p>
        """, 400


@app.route("/anniversaries")
def anniversaries():
    return evaluations()


@app.route("/birthdays-anniversaries", methods=["GET", "POST"])
def birthdays_anniversaries():
    if request.method == "GET":
        return render_template(
            "birthdays_anniversaries.html",
            default_month=datetime.now().strftime("%Y-%m"),
        )

    ukg_file = request.files.get("ukg_file")
    scope = request.form.get("scope", "month")
    report_month = request.form.get("report_month", "")
    if not ukg_file:
        return "<h2>Error</h2><p>The UKG report is required.</p>", 400
    if scope not in {"month", "year"}:
        return "<h2>Error</h2><p>Select the list needed.</p>", 400
    if scope == "month" and not report_month:
        return "<h2>Error</h2><p>Select a month.</p>", 400

    try:
        temp_folder = tempfile.mkdtemp()
        input_path = os.path.join(temp_folder, "ukg_birthdays_anniversaries.xlsx")
        ukg_file.save(input_path)
        employees, missing_birthdays = read_birthdays_and_anniversaries(input_path)
        if missing_birthdays:
            names = "".join(f"<li>{name}</li>" for name in missing_birthdays)
            return f"""
                <h2>Missing birthday information</h2>
                <p>The report stopped because these active employees have no birthday:</p>
                <ul>{names}</ul>
                <p>Correct the birthday information in UKG and run the report again.</p>
                <p><a href='/birthdays-anniversaries'>Try again</a></p>
            """, 400
        if not employees:
            raise ValueError("No active employees with birthday and anniversary dates were found.")

        month = None
        if scope == "month":
            selected = datetime.strptime(report_month, "%Y-%m")
            month = selected.month
            output_name = f"Birthdays_Anniversaries_{selected.strftime('%B')}.pdf"
        else:
            output_name = "Birthday_Anniversary_Company_List.pdf"

        output_path = os.path.join(temp_folder, output_name)
        create_birthday_anniversary_pdf(employees, output_path, month=month)
        return send_file(
            output_path,
            as_attachment=True,
            download_name=output_name,
            mimetype="application/pdf",
        )
    except Exception as error:
        return f"""
            <h2>Could not create the birthday and anniversary list</h2>
            <p>{error}</p>
            <p><a href='/birthdays-anniversaries'>Try again</a></p>
        """, 400


if __name__ == "__main__":
    app.run(debug=True)
